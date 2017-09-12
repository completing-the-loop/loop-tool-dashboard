import json

from datetime import datetime, time

from django.contrib.auth.decorators import login_required
from django.db import connections
from django.http import HttpResponse
from django.http import HttpResponseRedirect
from django.shortcuts import render_to_response
from django.shortcuts import redirect
from django.template import RequestContext
from docx import Document

from dashboard.models import CourseOffering
from dashboard.models import CourseRepeatingEvent
from dashboard.models import PedagogyHelper
from dashboard.utils import generate_userbycourse_histogram
from dashboard.utils import generate_userbyweek_histogram
from dashboard.utils import generate_usersforpage_histogram
from dashboard.utils import getSingleCourseEvents
from dashboard.utils import getSubmissionCourseEvents
from dashboard.utils import getusers_prepostevent_table
from dashboard.utils import get_avggrade
from dashboard.utils import get_contentdetails
from dashboard.utils import get_contentpageviews_dataset
from dashboard.utils import get_contenttypes
from dashboard.utils import get_courseweeks
from dashboard.utils import get_nocourseforumposts
from dashboard.utils import get_noforumposts
from dashboard.utils import get_prepostevent_table
from dashboard.utils import get_prepostevent_treetable
from dashboard.utils import get_quizattemps
from dashboard.utils import get_quizusercount
from dashboard.utils import get_quizusercoursecount
from dashboard.utils import get_userassessmentviews_dataset
from dashboard.utils import get_usercommunicationviews_dataset
from dashboard.utils import get_usercoursecount
from dashboard.utils import get_userdetails
from dashboard.utils import get_usergrade
from dashboard.utils import get_usernoforumposts
from dashboard.utils import get_userpageviews_dataset
from dashboard.utils import get_userquizattemps
from dashboard.utils import get_usersthatdidnotaccesscontent
from dashboard.utils import get_userweekcount
from dashboard.utils import weekbegend


@login_required
def pedagogyhelper(request):
    context = RequestContext(request)

    curuser = request.user

    course_id = None
    pedagogyhelper_json = '{}'
    save_pedagogyhelper_json = False

    if request.method == 'POST':
        course_id = int(request.POST['course_id'])
        pedagogyhelper_json = request.POST['pedagogyhelper_json']
        save_pedagogyhelper_json = True
    else:
        course_id = int(request.GET.get('course_id'))

    course_offering = CourseOffering.objects.get(pk=course_id)
    course_code = course_offering.code
    course_title = course_offering.title
    course_type = course_offering.lms_type

    if save_pedagogyhelper_json:
        pedagogyhelper_obj, created = PedagogyHelper.objects.update_or_create(course_offering=course_offering,defaults={'pedagogyhelper_json': pedagogyhelper_json})
        #pedagogyhelper_obj = PedagogyHelper(pedagogyhelper_json=pedagogyhelper_json, )
        #pedagogyhelper_obj.save()
    else:
        pedagogyhelper_objs = PedagogyHelper.objects.filter(course_offering=course_id)
        print("pedagogyhelper_objs", pedagogyhelper_objs)
        if (len(pedagogyhelper_objs) > 0):
            pedagogyhelper_json = pedagogyhelper_objs[0].pedagogyhelper_json
            print("from database", pedagogyhelper_json)

    cursor = connections['default'].cursor()

    sql = "SELECT sitetree FROM summary_courses WHERE course_id=%d" % (course_id)
    cursor.execute(sql);
    rows = cursor.fetchall()
    treelist_json = rows[0][0]
    #print(treelist_json)
    treelist  = json.loads(treelist_json)
    json_4_jstree = ""
    for courseitem in treelist:
        #print(courseitem)
        itemid = courseitem[0]
        parentid = courseitem[1]
        adjustedparentid = str(parentid)
        if parentid==0:
            adjustedparentid = "#"
        #get title and type
        res_sql = "SELECT title, content_type FROM dim_pages WHERE content_id=%d AND course_id=%d;" %(itemid, course_id)
        row_count = cursor.execute(res_sql);
        row = cursor.fetchall()
        title = "%s (%s)" % (row[0][0],row[0][1])
        json_4_jstree += '{ "id" : "%d", "parent" : "%s", "text" : "%s" },' % (itemid, adjustedparentid, title)

    context_dict = {'course_id':course_id, 'course_code': course_code, 'course_title':course_title, 'course_type':course_type, 'json_4_jstree':json_4_jstree, 'pedagogyhelper_json':pedagogyhelper_json}
    return render_to_response('dashboard/pedagogyhelper.html', context_dict, context)

def pedagogyhelperdownload(request):
    context = RequestContext(request)
    #PedagogyHelper.objects.all().delete()
    course_id = None
    pedagogyhelper_json = '{}'

    if request.method == 'POST':
        course_id = int(request.POST['course_id'])
        pedagogyhelper_json = request.POST['pedagogyhelper_exportjson']
        save_pedagogyhelper_json = True
    else:
        course_id = int(request.GET.get('course_id'))

    course_offering = CourseOffering.objects.get(pk=course_id)
    course_code = course_offering.code
    course_title = course_offering.title
    course_type = course_offering.lms_type

    if save_pedagogyhelper_json:
        pedagogyhelper_obj, created = PedagogyHelper.objects.update_or_create(course_offering=course_offering,defaults={'pedagogyhelper_json': pedagogyhelper_json})
        #pedagogyhelper_obj = PedagogyHelper(pedagogyhelper_json=pedagogyhelper_json, course_offering=course_offering)
        #pedagogyhelper_obj.save()

    pedagogyhelper_objs = PedagogyHelper.objects.filter(course_offering=course_id)

    pedagogyhelper_json = "{}"
    if (len(pedagogyhelper_objs) > 0):
        pedagogyhelper_json = pedagogyhelper_objs[0].pedagogyhelper_json

    document = Document()

    title = "%s: %s" % (course_code, course_title)

    document.add_heading(title, 0)

    if pedagogyhelper_json != "{}":

        pedagogyhelper  = json.loads(pedagogyhelper_json)

        table = document.add_table(rows=1, cols=3)
        table.style = 'TableGrid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Learning Objectives'
        hdr_cells[1].text = 'Learning Activities'
        hdr_cells[2].text = 'Learning Resources'
        for lo in pedagogyhelper['learningobjects']:
            row_cells = table.add_row().cells
            row_cells[0].text = "%s. %s" % (lo['id'], lo['name'])
            row_cells[1].text = ''
            row_cells[2].text = ''
            for la in lo['learningactivities']:
                row_cells = table.add_row().cells
                row_cells[0].text = ''
                row_cells[1].text = "%d. %s" % (int(la['id'])+1, la['name'])
                row_cells[2].text = ''
                for lr in la['learningresources']:
                    row_cells = table.add_row().cells
                    row_cells[0].text = ''
                    row_cells[1].text = ''
                    row_cells[2].text = "%d. %s" % (int(lr['id'])+1, lr['name'])

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=download.docx'
    document.save(response)

    return response

def coursedashboard(request):
    context = RequestContext(request)

    course_id = int(request.GET.get('course_id'))
    week_id = int(request.GET.get('week_filter'))

    if (week_id == -1):
        return redirect('/overallcoursedashboard?course_id=' + str(course_id))

    course_offering = CourseOffering.objects.get(pk=course_id)
    course_code = course_offering.code
    course_title = course_offering.title
    course_type = course_offering.lms_type
    if course_type == "Macquarie Uni Moodle":
        course_type = "MoodleMacquarie"

    course_weeks = get_courseweeks(course_id)

    if week_id == 0:
        week_id = course_weeks[0]

    week_no = course_weeks.index(week_id) + 1

    cursor = connections['default'].cursor()

    weekbeg, weekend, begweek_unix, endweek_unix = weekbegend(2015, week_id)

    #Pageview Graph - Content
    daytotal=[0,0,0,0,0,0,0]
    sql = "SELECT Date_dayinweek, Pageviews FROM Summary_CourseVisitsByDayInWeek WHERE DATE_WEEK = %d AND course_id=%d ORDER BY Date_dayinweek;" % (course_weeks[week_no-1], course_id)
    row_count = cursor.execute(sql);
    result = cursor.fetchall()
    for dayidx,row in enumerate(result):
        try:
             daytotal[int(row[0])] = int(row[1])
        except IndexError:
             print('Index Error')

    PageViewGraphContentList = ','.join(map(str, daytotal))

    #Pageview Graph - Communication
    daytotal=[0,0,0,0,0,0,0]
    sql = "SELECT Date_dayinweek, Pageviews FROM Summary_CourseCommunicationVisitsByDayInWeek WHERE DATE_WEEK = %d  AND course_id=%d ORDER BY Date_dayinweek;" % (course_weeks[week_no-1], course_id)
    row_count = cursor.execute(sql);
    result = cursor.fetchall()
    for dayidx,row in enumerate(result):
        try:
             daytotal[int(row[0])] = int(row[1])
        except IndexError:
             print('Index Error')

    PageViewGraphCommunicationList = ','.join(map(str, daytotal))

    #Pageview Graph - Assessment
    daytotal=[0,0,0,0,0,0,0]
    sql = "SELECT Date_dayinweek, Pageviews FROM Summary_CourseAssessmentVisitsByDayInWeek WHERE DATE_WEEK = %d  AND course_id=%d ORDER BY Date_dayinweek;" % (course_weeks[week_no-1], course_id)
    row_count = cursor.execute(sql);
    result = cursor.fetchall()
    for dayidx,row in enumerate(result):
        try:
             daytotal[int(row[0])] = int(row[1])
        except IndexError:
             print('Index Error')

    PageViewGraphAssessmentList = ','.join(map(str, daytotal))

    #Pageview Graph - Unique Views
    daytotal=[0,0,0,0,0,0,0]
    sql = "SELECT Date_dayinweek, Pageviews FROM Summary_UniquePageViewsByDayInWeek WHERE DATE_WEEK = %d  AND course_id=%d ORDER BY Date_dayinweek;" % (course_weeks[week_no-1], course_id)
    row_count = cursor.execute(sql);
    result = cursor.fetchall()
    for dayidx,row in enumerate(result):
        try:
             daytotal[int(row[0])] = int(row[1])
        except IndexError:
             print('Index Error')

    PageViewGraphUniqueViewsList = ','.join(map(str, daytotal))

    #Sparkline - Unique Pageviews
    daytotal=[0,0,0,0,0,0,0]
    sql = "SELECT Date_dayinweek, Pageviews FROM Summary_UniquePageViewsByDayInWeek WHERE DATE_WEEK = %d  AND course_id=%d ORDER BY Date_dayinweek;" % (course_weeks[week_no-1], course_id)
    row_count = cursor.execute(sql);
    result = cursor.fetchall()
    for dayidx,row in enumerate(result):
        try:
             daytotal[int(row[0])] = int(row[1])
        except IndexError:
             print('Index Error')

    uniquepageviewsbydayinweek = ','.join(map(str, daytotal))

    #Sparkline - Participating Users
    daytotal=[0,0,0,0,0,0,0]
    sql = "SELECT Date_dayinweek, Pageviews FROM Summary_ParticipatingUsersByDayInWeek WHERE DATE_WEEK = %d  AND course_id=%d ORDER BY Date_dayinweek;" % (course_weeks[week_no-1], course_id)
    row_count = cursor.execute(sql);
    result = cursor.fetchall()
    for dayidx,row in enumerate(result):
        try:
             daytotal[int(row[0])] = int(row[1])
        except IndexError:
             print('Index Error')

    participantsbydayinweek = ','.join(map(str, daytotal))

    #Sparkline - Sessions by Day in Week
    daytotal=[0,0,0,0,0,0,0]
    sql = "SELECT Date_dayinweek, sessions FROM Summary_SessionsByDayInWeek WHERE DATE_WEEK = %d AND course_id=%d ORDER BY Date_dayinweek;" % (course_weeks[week_no-1], course_id)
    row_count = cursor.execute(sql);
    result = cursor.fetchall()
    for dayidx,row in enumerate(result):
        try:
             daytotal[int(row[0])] = int(row[1])
        except IndexError:
             print('Index Error')

    sessionsbydayinweek = ','.join(map(str, daytotal))

    #Sparkline - Average Session Length by Day in Week
    daytotal=[0,0,0,0,0,0,0]
    sql = "SELECT Date_dayinweek, session_average_in_minutes FROM Summary_SessionAverageLengthByDayInWeek WHERE DATE_WEEK = %d  AND course_id=%d ORDER BY Date_dayinweek;" % (course_weeks[week_no-1], course_id)
    row_count = cursor.execute(sql);
    result = cursor.fetchall()
    for dayidx,row in enumerate(result):
        try:
             daytotal[int(row[0])] = int(row[1])
        except IndexError:
             print('Index Error')

    averagesessionlengthbydayinweek = ','.join(map(str, daytotal))

    #Sparkline - Average Pages Per Session by Day in Week
    daytotal=[0,0,0,0,0,0,0]
    sql = "SELECT Date_dayinweek, pages_per_session FROM Summary_SessionAveragePagesPerSessionByDayInWeek WHERE DATE_WEEK = %d  AND course_id=%d ORDER BY Date_dayinweek;" % (course_weeks[week_no-1], course_id)
    row_count = cursor.execute(sql);
    result = cursor.fetchall()
    for dayidx,row in enumerate(result):
        try:
             daytotal[int(row[0])] = int(row[1])
        except IndexError:
             print('Index Error')

    averagepagespersessionbydayinweek = ','.join(map(str, daytotal))

    histogram = generate_userbyweek_histogram(week_id, course_id)

    communication_types, assessment_types, communication_types_str, assessment_types_str = get_contenttypes(course_type)

    excluse_contentype_list = communication_types + assessment_types
    excluse_contentype_str = ','.join("'{0}'".format(x) for x in excluse_contentype_list)

    #Top Content
    sql = "SELECT SUM(F.pageview) AS Pageviews, F.page_id, P.title, F.module, F.action, F.url, P.order_no FROM dim_dates D LEFT JOIN fact_coursevisits F ON D.Id = F.Date_Id JOIN dim_pages P ON F.page_id=P.content_id  WHERE F.page_id!=0 AND F.module NOT IN (%s) AND D.Date_dayinweek IN (0,1,2,3,4,5,6) AND D.DATE_week IN (%s) AND F.course_id=%d GROUP BY F.page_id ORDER BY Pageviews DESC LIMIT 10;" % (excluse_contentype_str, course_weeks[week_no-1], course_id)
    cursor.execute(sql);
    topcontent_resultset = cursor.fetchall()
    topcontent_table = ""
    for row in topcontent_resultset:
        user_cnt = str(get_userweekcount(course_weeks[week_no-1],int(row[1]),  course_id, row[3], int(row[6])))
        topcontent_table = topcontent_table + '<tr><td><a href="/coursepage?page_id=%s&course_id=%s">%s</a></td><td class="center">%s</td><td class="center">%s</td><td class="center">%s</td></tr>' %(row[1],str(course_id), row[2],row[3],user_cnt,row[0])

    #Top Communcation Tools
    topforumcontent_table = ""
    sql = "SELECT SUM(F.pageview) AS Pageviews, F.page_id, P.title, F.module, F.action, F.url, P.order_no FROM dim_dates  D LEFT JOIN fact_coursevisits F ON D.Id = F.Date_Id JOIN dim_pages P ON F.page_id=P.content_id WHERE F.page_id!=0 AND F.module IN (%s) AND D.Date_dayinweek IN (0,1,2,3,4,5,6) AND D.DATE_week IN (%s) AND F.course_id=%d GROUP BY F.page_id ORDER BY Pageviews DESC LIMIT 10;" % (communication_types_str, course_weeks[week_no-1], course_id)
    cursor.execute(sql);
    topforumcontent_resultset = cursor.fetchall()
    topforumcontent_table = ""
    for row in topforumcontent_resultset:
        user_cnt = str(get_userweekcount(course_weeks[week_no-1],int(row[1]),  course_id, row[3], int(row[6])))
        no_posts = get_noforumposts(int(row[1]), course_id, course_weeks[week_no-1])
        #no_participants = getusercount(3,1)
        topforumcontent_table = topforumcontent_table + '<tr><td><a href="/coursepage?page_id=%s&course_id=%s">%s</a></td><td class="center">%s</td><td class="center">%s</td><td class="center">%s</td><td class="center">%d</td></tr>' %(row[1],str(course_id), row[2], row[3], user_cnt, row[0], no_posts)

    #Top Assessment Items
    sql = "SELECT SUM(F.pageview) AS Pageviews, F.page_id, P.title, F.module, F.action, F.url, P.order_no FROM dim_dates  D LEFT JOIN fact_coursevisits F ON D.Id = F.Date_Id JOIN dim_pages P ON F.page_id=P.content_id WHERE F.page_id!=0 AND F.module IN (%s) AND D.Date_dayinweek IN (0,1,2,3,4,5,6) AND D.DATE_week IN (%s) AND F.course_id=%d GROUP BY F.page_id ORDER BY Pageviews DESC LIMIT 10;" % (assessment_types_str, course_weeks[week_no-1], course_id)
    cursor.execute(sql);
    topquizcontent_resultset = cursor.fetchall()
    topquizcontent_table = ""
    for row in topquizcontent_resultset:
        user_cnt = str(get_quizusercount(course_weeks[week_no-1], int(row[1]), course_id)) #str(get_userweekcount(course_weeks[week_no-1],int(row[1]),  course_id, row[3], int(row[6])))
        attempts = str(get_quizattemps(int(row[1]), course_id, unixstart=begweek_unix, unixend=endweek_unix))
        averagestudentscore = str(get_avggrade(int(row[1]), course_id, unixstart=begweek_unix, unixend=endweek_unix))
        topquizcontent_table = topquizcontent_table + '<tr><td><a href="/coursepage?page_id=%s&course_id=%s">%s</a></td><td class="center">%s</td><td class="center">%s</td><td class="center">%s</td><td class="center">%s</td></tr>' %(row[1],str(course_id), row[2], row[3], user_cnt, attempts, averagestudentscore)

    #Top Users/Students
    sql = "SELECT COUNT(F.pageview) AS Pageviews, U.Firstname, U.Lastname, U.Role, F.user_id FROM dim_dates  D LEFT JOIN fact_coursevisits F ON D.Id = F.Date_Id JOIN dim_users U ON F.user_id=U.lms_user_id  WHERE D.DATE_week IN (%d) AND F.course_id=%d AND U.role='Student' GROUP BY F.user_id ORDER BY Pageviews DESC LIMIT 10;" % (course_weeks[week_no-1], course_id)
    cursor.execute(sql);
    topusers_resultset = cursor.fetchall()
    topusers_table = ""
    for row in topusers_resultset:
        topusers_table = topusers_table + '<tr><td class="center"><a href="/coursemember?user_id=%s_%s&course_id=%s">%s</a></td><td class="center">%s</td><td class="center">%s</td><td class="center">%s</td></tr>' %(str(course_id), row[4],str(course_id), row[1], row[2], row[3], row[0])

    #get repeating events
    repeating_evt = CourseRepeatingEvent.objects.filter(course_offering=course_offering)
    rpt_evt_lst = ""
    for evt in repeating_evt:
       rpt_evt_lst = rpt_evt_lst + "{ value: %s, color: 'green', width: 2, label: { text: '%s'}}," % (evt.day_of_week, evt.title)

    context_dict = {'repeatingevents': rpt_evt_lst, 'course_weeks': course_weeks, 'week_id': week_id, 'weekbeg': weekbeg, 'weekend': weekend, 'PageViewGraphContentList': PageViewGraphContentList, 'PageViewGraphCommunicationList': PageViewGraphCommunicationList, 'PageViewGraphAssessmentList': PageViewGraphAssessmentList, 'PageViewGraphUniqueViewsList': PageViewGraphUniqueViewsList, 'histogram_labels': histogram["labels"], 'histogram_values': histogram["values"], 'sessionsbydayinweek': averagepagespersessionbydayinweek, 'averagesessionlengthbydayinweek': averagesessionlengthbydayinweek, 'averagepagespersessionbydayinweek':averagepagespersessionbydayinweek, 'topusers_table': topusers_table, 'topcontent_table': topcontent_table, 'topforumcontent_table': topforumcontent_table,  'topquizcontent_table': topquizcontent_table, 'course_id': course_id, 'course_code': course_code, 'course_title': course_title, 'week_no': week_no, 'uniquepageviewsbydayinweek': uniquepageviewsbydayinweek, 'participantsbydayinweek': participantsbydayinweek}

    return render_to_response('dashboard/coursedashboard.html', context_dict, context)

@login_required
def overallcoursedashboard(request):
    context = RequestContext(request)

    cursor = connections['default'].cursor()

    course_id = int(request.GET.get('course_id'))

    course_offering = CourseOffering.objects.get(pk=course_id)
    course_code = course_offering.code
    course_title = course_offering.title
    course_type = course_offering.lms_type
    if course_type == "Macquarie Uni Moodle":
        course_type = "MoodleMacquarie"

    course_weeks = get_courseweeks(course_id)

    weeks = ','.join(map(str, course_weeks))

    communication_types, assessment_types, communication_types_str, assessment_types_str = get_contenttypes(course_type)

    excluse_contentype_list = communication_types + assessment_types
    excluse_contentype_str = ','.join("'{0}'".format(x) for x in excluse_contentype_list)

    cursor = connections['default'].cursor()

    histogram = generate_userbycourse_histogram(course_id)

    sql = "SELECT contentcoursepageviews, communicationcoursepageviews, assessmentcoursepageviews FROM summary_courses WHERE course_id=%d" % (course_id)
    cursor.execute(sql);
    rows = cursor.fetchall()
    contentviews =  rows[0][0]
    communicationviews =  rows[0][1]
    assessmentviews =  rows[0][2]

    #Top Content
    sql = "SELECT SUM(F.pageview) AS Pageviews, F.page_id, P.title, F.module, F.action, F.url, P.order_no FROM dim_dates  D LEFT JOIN fact_coursevisits F ON D.Id = F.Date_Id JOIN dim_pages P ON F.page_id=P.content_id  WHERE F.page_id!=0 AND F.module NOT IN (%s) AND F.course_id=%d GROUP BY F.page_id ORDER BY Pageviews DESC LIMIT 10;" % (excluse_contentype_str, course_id)
    cursor.execute(sql);
    topcontent_resultset = cursor.fetchall()
    topcontent_table = ""
    for row in topcontent_resultset:
        user_cnt = str(get_usercoursecount(int(row[1]),  course_id, row[3], int(row[6])))
        topcontent_table = topcontent_table + '<tr><td><a href="/coursepage?page_id=%s&course_id=%s">%s</a></td><td class="center">%s</td><td class="center">%s</td><td class="center">%s</td></tr>' %(row[1],str(course_id), row[2],row[3],user_cnt,row[0])

    #Top Communcation Tools
    topforumcontent_table = ""
    sql = "SELECT SUM(F.pageview) AS Pageviews, F.page_id, P.title, F.module, F.action, F.url, P.order_no FROM dim_dates  D LEFT JOIN fact_coursevisits F ON D.Id = F.Date_Id JOIN dim_pages P ON F.page_id=P.content_id WHERE F.page_id!=0 AND F.module IN (%s) AND F.course_id=%d GROUP BY F.page_id ORDER BY Pageviews DESC LIMIT 10;" % (communication_types_str, course_id)
    cursor.execute(sql);
    topforumcontent_resultset = cursor.fetchall()
    topforumcontent_table = ""
    for row in topforumcontent_resultset:
        user_cnt = str(get_usercoursecount(int(row[1]),  course_id, row[3], int(row[6])))
        no_posts = get_nocourseforumposts(int(row[1]), course_id)
        #no_participants = getusercount(3,1)
        topforumcontent_table = topforumcontent_table + '<tr><td><a href="/coursepage?page_id=%s&course_id=%s">%s</a></td><td class="center">%s</td><td class="center">%s</td><td class="center">%s</td><td class="center">%d</td></tr>' %(row[1],str(course_id), row[2], row[3], user_cnt, row[0], no_posts)

    #Top Assessment Items
    sql = "SELECT SUM(F.pageview) AS Pageviews, F.page_id, P.title, F.module, F.action, F.url, P.order_no FROM dim_dates  D LEFT JOIN fact_coursevisits F ON D.Id = F.Date_Id JOIN dim_pages P ON F.page_id=P.content_id WHERE F.page_id!=0 AND F.module IN (%s) AND F.course_id=%d GROUP BY F.page_id ORDER BY Pageviews DESC LIMIT 10;" % (assessment_types_str, course_id)
    cursor.execute(sql);
    topquizcontent_resultset = cursor.fetchall()
    topquizcontent_table = ""
    for row in topquizcontent_resultset:
        user_cnt = str(get_quizusercoursecount(int(row[1]),  course_id)) #str(get_usercoursecount(int(row[1]),  course_id, row[3], int(row[6])))
        attempts = str(get_quizattemps(int(row[1]), course_id))
        averagestudentscore = str(get_avggrade(int(row[1]), course_id))
        topquizcontent_table = topquizcontent_table + '<tr><td><a href="/coursepage?page_id=%s&course_id=%s">%s</a></td><td class="center">%s</td><td class="center">%s</td><td class="center">%s</td><td class="center">%s</td></tr>' %(row[1],str(course_id), row[2], row[3], user_cnt, attempts, averagestudentscore)
        #topquizcontent_table = topquizcontent_table + '<tr><td><a href="/coursepage?page_id=%s&course_id=%s">%s</a></td><td class="center">%s</td><td class="center">%s</td><td class="center">%s</td><td class="center">%s</td><td class="center">%s</td></tr>' %(row[1],str(course_id), row[2], row[3], user_cnt, row[0], attempts, averagestudentscore)

    #Top Users/Students
    sql = "SELECT COUNT(F.pageview) AS Pageviews, U.Firstname, U.Lastname, U.Role, F.user_id FROM dim_dates  D LEFT JOIN fact_coursevisits F ON D.Id = F.Date_Id JOIN dim_users U ON F.user_id=U.lms_user_id  WHERE F.course_id=%d GROUP BY F.user_id ORDER BY Pageviews DESC LIMIT 10;" % (course_id)
    cursor.execute(sql);
    topusers_resultset = cursor.fetchall()
    topusers_table = ""
    for row in topusers_resultset:
        topusers_table = topusers_table + '<tr><td class="center"><a href="/coursemember?user_id=%s&course_id=%s">%s</a></td><td class="center">%s</td><td class="center">%s</td><td class="center">%s</td></tr>' %(row[4],str(course_id), row[1], row[2], row[3], row[0])

    #get Event flags
    sng_evt_flags = getSingleCourseEvents(course_id)
    sub_evt_flags = getSubmissionCourseEvents(course_id)

    context_dict = {'sng_evt_flags': sng_evt_flags, 'sub_evt_flags':sub_evt_flags, 'contentviews': contentviews, 'communicationviews': communicationviews, 'assessmentviews':assessmentviews, 'course_weeks': course_weeks, 'week_id': -1, 'histogram_labels': histogram["labels"], 'histogram_values': histogram["values"], 'topusers_table': topusers_table, 'topcontent_table': topcontent_table, 'topforumcontent_table': topforumcontent_table,  'topquizcontent_table': topquizcontent_table, 'course_id': course_id, 'course_code': course_code, 'course_title': course_title}

    return render_to_response('dashboard/overallcoursedashboard.html', context_dict, context)

@login_required
def coursemembers(request):
    context = RequestContext(request)

    curr_evt = None
    course_id = None

    context = RequestContext(request)
    if request.method == 'POST':
        course_id = int(request.POST['course_id'])
        curr_evt = int(request.POST['repevt'])
    else:
        course_id = int(request.GET.get('course_id'))

    course_weeks = get_courseweeks(course_id)

    weeks = ','.join(map(str, course_weeks))

    course_offering = CourseOffering.objects.get(pk=course_id)
    course_code = course_offering.code
    course_title = course_offering.title
    course_type = course_offering.lms_type
    if course_type == "Macquarie Uni Moodle":
        course_type = "MoodleMacquarie"

    cursor = connections['default'].cursor()

    #Get Counts and Vis Table # users_vis_table
    sql = "SELECT users_counts_table FROM summary_courses WHERE course_id=%d" % (course_id)
    cursor.execute(sql);
    rows = cursor.fetchall()
    users_counts_table =  rows[0][0]

    #get repeating events
    weekdays = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
    repeating_evt = CourseRepeatingEvent.objects.filter(course_offering=course_offering)
    opts = ""
    count = 0
    first_evt_day = None
    for evt in repeating_evt:
        if count==0: first_evt_day=int(evt.day_of_week)
        selected = ""
        if curr_evt==int(evt.day_of_week):
            selected = "selected"
        opts = opts + "<option value='%s' %s>%s (%s)</option>" % (evt.day_of_week, selected, evt.title, weekdays[int(evt.day_of_week)])
        count = count + 1

    users_vis_table = ""
    showeventtab = False
    if curr_evt is not None:
        showeventtab = True
        users_vis_table = getusers_prepostevent_table(course_id, curr_evt, course_weeks)
    else:
        users_vis_table = getusers_prepostevent_table(course_id, first_evt_day, course_weeks)

    context_dict = {'showeventtab': showeventtab, 'opts': opts, 'repeating_evt':repeating_evt, 'course_id': course_id, 'course_code': course_code, 'course_title': course_title, 'users_counts_table': users_counts_table, 'users_vis_table': users_vis_table}

    return render_to_response('dashboard/coursemembers.html', context_dict, context)


@login_required
def coursestructure(request):
    context = RequestContext(request)

    course_id = int(request.GET.get('course_id'))

    course_offering = CourseOffering.objects.get(pk=course_id)
    course_code = course_offering.code
    course_title = course_offering.title
    course_type = course_offering.lms_type
    if course_type == "Macquarie Uni Moodle":
        course_type = "MoodleMacquarie"

    cursor = connections['default'].cursor()

    #Get Counts and Vis Table
    sql = "SELECT access_counts_table, access_vis_table FROM summary_courses WHERE course_id=%d" % (course_id)
    cursor.execute(sql);
    rows = cursor.fetchall()
    access_counts_table =  rows[0][0]
    access_vis_table =  rows[0][1]

    repeating_evt = CourseRepeatingEvent.objects.filter(course_offering=course_offering)

    context_dict = {'repeating_evt': repeating_evt, 'course_id': course_id, 'course_code': course_code, 'course_title': course_title, 'access_counts_table': access_counts_table, 'access_vis_table': access_vis_table}

    return render_to_response('dashboard/coursestructure.html', context_dict, context)

@login_required
def coursemember(request):
    context = RequestContext(request)
    course_id = None
    member_id = None

    startdte = None
    enddte = None
    unixstart = None
    unixend = None
    startUTC = ""
    endUTC = ""
    if request.method == 'POST':
        course_id = int(request.POST['course_id'])
        member_id = request.POST['user_id']
        startdte = request.POST['start']
        enddte = request.POST['end']
    else:
        member_id = request.GET.get('user_id')
        course_id = int(request.GET.get('course_id'))

    range_str = ""
    if (startdte is not None):
        range_str = "(" + startdte + " - " + enddte + ")"
        start = datetime.strptime(startdte, "%Y-%m-%d")
        end = datetime.strptime(enddte, "%Y-%m-%d")
        startUTC = 'Date.UTC(%d, %d, %d)' % (start.year, start.month-1, start.day)
        endUTC = 'Date.UTC(%d, %d, %d)' % (end.year, end.month-1, end.day)
        unixstart = time.mktime(datetime.strptime(startdte, "%Y-%m-%d").timetuple())
        unixend = time.mktime(datetime.strptime(enddte, "%Y-%m-%d").timetuple())

    course_offering = CourseOffering.objects.get(pk=course_id)
    course_code = course_offering.code
    course_title = course_offering.title
    course_type = course_offering.lms_type
    if course_type == "Macquarie Uni Moodle":
        course_type = "MoodleMacquarie"

    userdetails = get_userdetails(member_id, course_id)
    firstname =  userdetails['firstname']
    surname = userdetails['lastname']
    email = userdetails['email']
    role = userdetails['role']

    course_weeks = get_courseweeks(course_id)
    weeks = ','.join(map(str, course_weeks))

    contentviews = get_userpageviews_dataset(member_id, course_id, weeks, course_type)
    assessmentviews = get_userassessmentviews_dataset(member_id, course_id, weeks, course_type)
    communicationsviews = get_usercommunicationviews_dataset(member_id, course_id, weeks, course_type)

    user_id = int(member_id[2:(len(member_id))])

    communication_types, assessment_types, communication_types_str, assessment_types_str = get_contenttypes(course_type)

    excluse_contentype_list = communication_types + assessment_types
    excluse_contentype_str = ','.join("'{0}'".format(x) for x in excluse_contentype_list)

    cursor = connections['default'].cursor()

    #Communication Tools - User Pageviews and No Posts
    if (startdte is None):
        sql = "SELECT SUM(F.pageview) AS Pageviews, F.page_id, P.title, F.module, F.action, F.url, P.order_no FROM dim_dates  D LEFT JOIN fact_coursevisits F ON D.Id = F.Date_Id JOIN dim_pages P ON F.page_id=P.content_id WHERE F.page_id!=0 AND F.module IN (%s) AND F.course_id=%d AND F.user_id=%d GROUP BY F.page_id ORDER BY Pageviews;" % (communication_types_str, course_id, user_id)
    else:
        sql = "SELECT SUM(F.pageview) AS Pageviews, F.page_id, P.title, F.module, F.action, F.url, P.order_no FROM dim_dates  D LEFT JOIN fact_coursevisits F ON D.Id = F.Date_Id JOIN dim_pages P ON F.page_id=P.content_id WHERE F.page_id!=0 AND F.module IN (%s) AND F.course_id=%d AND F.user_id=%d AND F.unixtimestamp BETWEEN %d AND %d GROUP BY F.page_id ORDER BY Pageviews;" % (communication_types_str, course_id, user_id, unixstart, unixend)
    print(sql)
    cursor.execute(sql);
    topforumcontent_resultset = cursor.fetchall()
    topforumcontent_table = ""
    for row in topforumcontent_resultset:
        no_posts = get_usernoforumposts(int(row[1]), course_id, user_id, unixstart=unixstart, unixend=unixend)
        topforumcontent_table = topforumcontent_table + '<tr><td>%s</td><td class="center">%s</td><td class="center">%s</td><td class="center">%d</td></tr>' %(row[2], row[3], row[0], no_posts)

    #Assessment Items - User Pageviews and Attempts
    if (startdte is None):
        sql = "SELECT SUM(F.pageview) AS Pageviews, F.page_id, P.title, F.module, F.action, F.url, P.order_no FROM dim_dates  D LEFT JOIN fact_coursevisits F ON D.Id = F.Date_Id JOIN dim_pages P ON F.page_id=P.content_id WHERE F.page_id!=0 AND F.module IN (%s) AND F.course_id=%d AND F.user_id=%d  GROUP BY F.page_id ORDER BY Pageviews;" % (assessment_types_str, course_id, user_id)
    else:
        sql = "SELECT SUM(F.pageview) AS Pageviews, F.page_id, P.title, F.module, F.action, F.url, P.order_no FROM dim_dates  D LEFT JOIN fact_coursevisits F ON D.Id = F.Date_Id JOIN dim_pages P ON F.page_id=P.content_id WHERE F.page_id!=0 AND F.module IN (%s) AND F.course_id=%d AND F.user_id=%d AND F.unixtimestamp BETWEEN %d AND %d  GROUP BY F.page_id ORDER BY Pageviews;" % (assessment_types_str, course_id, user_id, unixstart, unixend)
    print(sql)
    cursor.execute(sql);
    topquizcontent_resultset = cursor.fetchall()
    topquizcontent_table = ""
    for row in topquizcontent_resultset:
        attempts = str(get_userquizattemps(int(row[1]), course_id, user_id, unixstart=unixstart, unixend=unixend))
        studentscore = str(get_usergrade(int(row[1]), course_id, user_id, unixstart=unixstart, unixend=unixend))
        topquizcontent_table = topquizcontent_table + '<tr><td>%s</td><td class="center">%s</td><td class="center">%s</td><td class="center">%s</td><td class="center">%s</td></tr>' %(row[2], row[3], row[0], attempts, studentscore)

    #get Event flags
    sng_evt_flags = getSingleCourseEvents(course_id)
    sub_evt_flags = getSubmissionCourseEvents(course_id)

    context_dict = {'range_str': range_str, 'startUTC': startUTC, 'endUTC': endUTC, 'sng_evt_flags': sng_evt_flags, 'sub_evt_flags':sub_evt_flags, 'topforumcontent_table': topforumcontent_table, 'topquizcontent_table': topquizcontent_table, 'contentviews': contentviews, 'communicationsviews':communicationsviews, 'assessmentviews':assessmentviews, 'course_id': course_id, 'course_code': course_code, 'course_title': course_title, 'member_id':member_id, 'user_id': user_id, 'firstname':firstname, 'surname':surname, 'email':email, 'role':role}

    return render_to_response('dashboard/coursemember.html', context_dict, context)

@login_required
def coursepage(request):
    context = RequestContext(request)
    course_id = None
    page_id = None

    startdte = None
    enddte = None
    unixstart = None
    unixend = None
    startUTC = ""
    endUTC = ""
    if request.method == 'POST':
        course_id = int(request.POST['course_id'])
        page_id = int(request.POST['page_id'])
        startdte = request.POST['start']
        enddte = request.POST['end']
    else:
        page_id = int(request.GET.get('page_id'))
        course_id = int(request.GET.get('course_id'))

    section_order = 0

    course_offering = CourseOffering.objects.get(pk=course_id)
    course_code = course_offering.code
    course_title = course_offering.title
    course_type = course_offering.lms_type
    if course_type == "Macquarie Uni Moodle":
        course_type = "MoodleMacquarie"

    contentdetails = get_contentdetails(page_id, course_id)
    title =  contentdetails["title"]
    content_type = contentdetails["type"]

    if (content_type=="section"):
        if request.method == 'POST':
            section_order = int(request.POST['section_order'])
        else:
            section_order = int(request.GET.get('section_order'))

    range_str = ""
    if (startdte is not None):
        range_str = "(" + startdte + " - " + enddte + ")"
        start = datetime.strptime(startdte, "%Y-%m-%d")
        end = datetime.strptime(enddte, "%Y-%m-%d")
        startUTC = 'Date.UTC(%d, %d, %d)' % (start.year, start.month-1, start.day)
        endUTC = 'Date.UTC(%d, %d, %d)' % (end.year, end.month-1, end.day)
        unixstart = time.mktime(datetime.strptime(startdte, "%Y-%m-%d").timetuple())
        unixend = time.mktime(datetime.strptime(enddte, "%Y-%m-%d").timetuple())

    course_weeks = get_courseweeks(course_id)

    weeks = ','.join(map(str, course_weeks))

    communication_types, assessment_types, communication_types_str, assessment_types_str = get_contenttypes(course_type)

    excluse_contentype_list = communication_types + assessment_types
    excluse_contentype_str = ','.join("'{0}'".format(x) for x in excluse_contentype_list)

    # get content_type and then decide what
    contentviews = get_contentpageviews_dataset(page_id, course_id, weeks,section_order, course_type)

    histogram = generate_usersforpage_histogram(course_id, page_id, section_order, unixstart=unixstart, unixend=unixend)

    user_noaccess_table = get_usersthatdidnotaccesscontent(page_id, course_id, section_order, unixstart=unixstart, unixend=unixend)
    no_posts = ""
    attempts = ""
    averagestudentscore = ""
    if content_type in communication_types:
        no_posts = str(get_nocourseforumposts(page_id, course_id, unixstart=unixstart, unixend=unixend))
    elif content_type in assessment_types:
        attempts = str(get_quizattemps(page_id, course_id, unixstart=unixstart, unixend=unixend))
        averagestudentscore = str(get_avggrade(page_id, course_id, unixstart=unixstart, unixend=unixend))

    #get Event flags
    sng_evt_flags = getSingleCourseEvents(course_id)
    sub_evt_flags = getSubmissionCourseEvents(course_id)

    context_dict = {'range_str': range_str, 'startUTC': startUTC, 'endUTC': endUTC, 'sng_evt_flags': sng_evt_flags, 'sub_evt_flags':sub_evt_flags, 'no_posts': no_posts, 'attempts': attempts, 'averagestudentscore': averagestudentscore, 'histogram_labels': histogram["labels"], 'histogram_values': histogram["values"], 'user_noaccess_table': user_noaccess_table, 'contentviews': contentviews, 'course_id': course_id, 'course_code': course_code, 'course_title': course_title, 'page_id':page_id, 'title':title, 'content_type':content_type, 'section_order': section_order, 'communication_types':communication_types, 'assessment_types': assessment_types }

    return render_to_response('dashboard/coursepage.html',context_dict , context)

@login_required
def content(request):
    curr_evt = None
    course_id = None

    context = RequestContext(request)
    if request.method == 'POST':
        course_id = int(request.POST['course_id'])
        curr_evt = int(request.POST['repevt'])
    else:
        course_id = int(request.GET.get('course_id'))

    course_weeks = get_courseweeks(course_id)

    weeks = ','.join(map(str, course_weeks))

    course_offering = CourseOffering.objects.get(pk=course_id)
    course_code = course_offering.code
    course_title = course_offering.title
    course_type = course_offering.lms_type
    if course_type == "Macquarie Uni Moodle":
        course_type = "MoodleMacquarie"

    showeventtab = False
    if curr_evt is not None:
        showeventtab = True

    cursor = connections['default'].cursor()

    sql = "SELECT content_counts_table, content_user_table, sitetree FROM summary_courses WHERE course_id=%d" % (course_id)
    cursor.execute(sql);
    rows = cursor.fetchall()
    content_counts_table =  rows[0][0]
    content_user_table =  rows[0][1]
    treelist_json = rows[0][2]

    #get repeating events
    weekdays = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
    repeating_evt = CourseRepeatingEvent.objects.filter(course_offering=course_offering)
    opts = ""
    count = 0
    first_evt_day = None
    for evt in repeating_evt:
        if count==0: first_evt_day=int(evt.day_of_week)
        selected = ""
        if curr_evt==int(evt.day_of_week):
            selected = "selected"
        opts = opts + "<option value='%s' %s>%s (%s)</option>" % (evt.day_of_week, selected, evt.title, weekdays[int(evt.day_of_week)])
        count = count + 1

    vis_table = ""
    showeventtab = False
    if curr_evt is not None:
        showeventtab = True
        vis_table = get_prepostevent_treetable(course_id, course_weeks, treelist_json, curr_evt, course_type)
    else:
        vis_table = get_prepostevent_treetable(course_id, course_weeks, treelist_json, first_evt_day, course_type)

    context_dict = {'vis_table': vis_table, 'showeventtab': showeventtab, 'opts': opts, 'repeating_evt':repeating_evt,'course_id': course_id, 'course_code': course_code, 'course_title': course_title, 'content_counts_table': content_counts_table, 'content_user_table': content_user_table}

    return render_to_response('dashboard/content.html',context_dict , context)

@login_required
def communication(request):

    curr_evt = None
    course_id = None

    context = RequestContext(request)
    if request.method == 'POST':
        course_id = int(request.POST['course_id'])
        curr_evt = int(request.POST['repevt'])
    else:
        course_id = int(request.GET.get('course_id'))

    context = RequestContext(request)
    course_weeks = get_courseweeks(course_id)

    weeks = ','.join(map(str, course_weeks))

    course_offering = CourseOffering.objects.get(pk=course_id)
    course_code = course_offering.code
    course_title = course_offering.title
    course_type = course_offering.lms_type
    if course_type == "Macquarie Uni Moodle":
        course_type = "MoodleMacquarie"

    cursor = connections['default'].cursor()

    sql = "SELECT communication_counts_table, communication_user_table, forum_posts_table FROM summary_courses WHERE course_id=%d" % (course_id)
    cursor.execute(sql);
    rows = cursor.fetchall()
    communication_counts_table = rows[0][1]
    communication_user_table = rows[0][0]
    forum_posts_table = rows[0][2]

    #get repeating events
    weekdays = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
    repeating_evt = CourseRepeatingEvent.objects.filter(course_offering=course_offering)
    opts = ""
    count = 0
    first_evt_day = None
    for evt in repeating_evt:
        if count==0: first_evt_day=int(evt.day_of_week)
        selected = ""
        if curr_evt==int(evt.day_of_week):
            selected = "selected"
        opts = opts + "<option value='%s' %s>%s (%s)</option>" % (evt.day_of_week, selected, evt.title, weekdays[int(evt.day_of_week)])
        count = count + 1

    vis_table = ""
    showeventtab = False
    if curr_evt is not None:
        showeventtab = True
        vis_table = get_prepostevent_table(course_id, curr_evt, 'forum', course_weeks, course_type)
    else:
        vis_table = get_prepostevent_table(course_id, first_evt_day, 'forum', course_weeks, course_type)

    context_dict = {'vis_table': vis_table, 'showeventtab': showeventtab, 'opts': opts, 'repeating_evt':repeating_evt, 'course_id': course_id, 'course_code': course_code, 'course_title': course_title, 'communication_counts_table': communication_counts_table, 'communication_user_table': communication_user_table, 'forum_posts_table': forum_posts_table}

    return render_to_response('dashboard/communication.html',context_dict , context)

@login_required
def assessment(request):
    context = RequestContext(request)

    curr_evt = None
    course_id = None

    context = RequestContext(request)
    if request.method == 'POST':
        course_id = int(request.POST['course_id'])
        curr_evt = int(request.POST['repevt'])
    else:
        course_id = int(request.GET.get('course_id'))

    course_weeks = get_courseweeks(course_id)

    weeks = ','.join(map(str, course_weeks))

    course_offering = CourseOffering.objects.get(pk=course_id)
    course_code = course_offering.code
    course_title = course_offering.title
    course_type = course_offering.lms_type
    if course_type == "Macquarie Uni Moodle":
        course_type = "MoodleMacquarie"

    cursor = connections['default'].cursor()

    sql = "SELECT assessment_counts_table, assessment_user_table, assessmentgrades FROM summary_courses WHERE course_id=%d" % (course_id)
    cursor.execute(sql);
    rows = cursor.fetchall()
    assessment_counts_table = rows[0][1]
    assessment_user_table = rows[0][0]
    assessment_grades_table = rows[0][2]

    #get repeating events
    weekdays = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
    repeating_evt = CourseRepeatingEvent.objects.filter(course_offering=course_offering)
    opts = ""
    count = 0
    first_evt_day = None
    for evt in repeating_evt:
        if count==0: first_evt_day=int(evt.day_of_week)
        selected = ""
        if curr_evt==int(evt.day_of_week):
            selected = "selected"
        opts = opts + "<option value='%s' %s>%s (%s)</option>" % (evt.day_of_week, selected, evt.title, weekdays[int(evt.day_of_week)])
        count = count + 1

    vis_table = ""
    showeventtab = False
    if curr_evt is not None:
        showeventtab = True
        vis_table = get_prepostevent_table(course_id, curr_evt, 'assign', course_weeks, course_type)
    else:
        vis_table = get_prepostevent_table(course_id, first_evt_day, 'assign', course_weeks, course_type)

    context_dict = {'vis_table': vis_table, 'showeventtab': showeventtab, 'opts': opts, 'repeating_evt':repeating_evt, 'course_id': course_id, 'course_code': course_code, 'course_title': course_title, 'assessment_counts_table': assessment_counts_table, 'assessment_user_table': assessment_user_table, 'assessment_grades_table': assessment_grades_table}

    return render_to_response('dashboard/assessment.html',context_dict , context)

@login_required
def loggedout(request):
    logout(request)
    # Take the user back to the homepage.
    return HttpResponseRedirect('/')
